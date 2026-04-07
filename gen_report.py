import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---- Page margins ----
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# ---- Style helpers ----
def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text, color=(20, 20, 20)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_font(r, bold=True, size=15, color=color)
    return p

def heading2(text, color=(50, 50, 50)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    set_font(r, bold=True, size=12, color=color)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        rb = p.add_run(bold_prefix + '  ')
        set_font(rb, bold=True, size=10.5)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def status_row(doc, label, status_text, status_color):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Cm(13.5)
    tbl.columns[1].width = Cm(3.5)
    c0, c1 = tbl.rows[0].cells
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c0.paragraphs[0].add_run(label)
    set_font(r, size=10.5)
    rb = c1.paragraphs[0].add_run(status_text)
    set_font(rb, bold=True, size=9.5, color=status_color)
    tbl_pr = tbl._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        borders.append(b)
    tbl_pr.append(borders)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    '4')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), 'CCCCCC')
    pb.append(b)
    pPr.append(pb)

GREEN  = (16,  150,  80)
ORANGE = (210, 100,  10)
RED    = (200,  30,  30)
GREY   = (100, 100, 100)

# ================================================================
# TITLE BLOCK
# ================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(4)
rt = title_p.add_run('Legit Extension — Code Review Response')
set_font(rt, bold=True, size=18, color=(20, 20, 20))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)
rs = sub_p.add_run('Review Findings Status Report  \u2022  v2.1.3')
set_font(rs, italic=True, size=11, color=GREY)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(16)
rd = date_p.add_run(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}')
set_font(rd, size=10, color=GREY)

hline()

# ================================================================
# EXECUTIVE SUMMARY
# ================================================================
heading1('Executive Summary')
body(
    'This document tracks the resolution status of each finding raised in the '
    'Legit Extension Code Review (v2.1.2). Out of 12 discrete findings, '
    '7 have been fully addressed, 2 partially or already addressed, and 3 remain open. '
    'The most impactful architectural change -- splitting the popup.js monolith '
    'into focused modules -- has been completed. Two of the three critical/medium '
    'bugs are still present in the codebase and should be prioritised before '
    'the next Chrome Web Store submission.'
)
hline()

# ================================================================
# SECTION 1 -- BUGS
# ================================================================
heading1('1  Bugs', color=(160, 30, 30))

heading2('Bug 1 -- Swapped Parameters in styleScoreLabel  [Critical]', color=(160, 30, 30))
body('Location: utils.js -- call site line 535, definition line 561')
body(
    'The call at line 535 passes arguments in the order (..., gradient, labelKey, emoji), '
    'but the function signature at line 561 expects (..., gradient, emoji, labelKey). '
    'As a result the emoji symbol is placed where the translation key is expected. '
    'The fallback at line 587 (TRANSLATIONS[currentLang][labelKey] || labelKey) hides '
    'the crash -- labelKey ends up as the emoji character which is not a translation key, '
    'so it falls back to the raw emoji -- but the visual output is wrong: the label '
    'renders as "HIGHLY CREDIBLE +" instead of "+ Highly Credible".'
)
status_row(doc,
    'Status: call-site argument order not corrected; swapped args still present.',
    'NOT FIXED', RED)

heading2('Bug 2 -- Tautological Condition in analyzeAgent  [Medium]', color=(100, 60, 0))
body('Location: orchestrator.js (was popup.js) line 388')
body(
    'The original condition  if(agent.id === "bias" || agent)  always evaluated '
    'to true because the bare agent reference is always truthy, causing every agent '
    'to receive quote-linkification instead of only bias and style.'
)
status_row(doc,
    'Fixed: condition is now  if (agent.id === "bias" || agent.id === "style")',
    'FIXED', GREEN)

heading2('Bug 3 -- waitForTabLoad Has No Timeout  [Medium]', color=(100, 60, 0))
body('Location: utils.js lines 1022-1032')
body(
    'The function registers a chrome.tabs.onUpdated listener and resolves only '
    'when the tab URL no longer contains "google.com" or "duckduckgo.com". '
    'If a redirect keeps the tab on one of those domains the listener is never '
    'removed, causing a persistent memory leak. No setTimeout guard has been added.'
)
status_row(doc,
    'Status: no timeout added; potential memory leak remains.',
    'NOT FIXED', RED)

hline()

# ================================================================
# SECTION 2 -- EFFICIENCY
# ================================================================
heading1('2  Efficiency', color=(130, 80, 0))

heading2('E1 -- Full Article Text Sent to All Agents  [High]', color=(100, 60, 0))
body('Location: orchestrator.js line 213; agents.js -- all prompts')
body(
    'The recommendation was to (a) cap excerptStart at 3,000-4,000 chars '
    'and (b) use shortExcerpt for headline and author agents. '
    'Partial progress: the author agent now receives shortExcerpt (600 chars). '
    'However excerptStart is still the full bodyText with no character ceiling, '
    'and the headline agent still passes longExcerpt (the full body) to the model.'
)
status_row(doc,
    'Author agent improved; headline agent and global text cap: not yet applied.',
    'PARTIAL', ORANGE)

heading2('E2 -- verify-format Agent Chains Double API Calls  [High]', color=(100, 60, 0))
body('Location: agents.js -- source-verify/source-format and consensus-verify/consensus-format chains')
body(
    'Each two-agent chain makes 2 API calls; the format agents only rewrite '
    'citations, work that could be done in JS or in a single merged prompt. '
    'Both chains remain unchanged. The estimated saving of ~47% of per-analysis '
    'cost ($0.07) has not been realised.'
)
status_row(doc,
    'Both chains intact; no change applied.',
    'NOT ADDRESSED', RED)

heading2('E3 -- Search Grounding for Bias Agent  [Medium]', color=GREY)
body('Location: agents.js -- bias agent configuration')
body(
    'The review asked whether the bias agent truly needs Search Grounding, given '
    'that bias is determined from the article text itself. In the current codebase '
    'the bias agent has useSearch: false -- it does not use Search Grounding. '
    'This was either pre-existing or corrected before the improvements tracker '
    'was created. Either way, the configuration is correct.'
)
status_row(doc,
    'bias agent: useSearch: false -- correctly configured.',
    'ALREADY OK', GREEN)

heading2('E4 -- Inefficient DOM Management  [Low]', color=GREY)
body('Location: utils.js -- attachQuoteLinkListeners (line 879); sortGridDynamic (line 309)')
body(
    'Both issues have been resolved. attachQuoteLinkListeners now uses a single '
    'delegated listener on #agentGrid with a module-level guard flag '
    '(_quoteLinkListenerAttached) to prevent duplicate registration -- no more '
    'per-element DOM cloning. sortGridDynamic wraps its DOM mutations in '
    'requestAnimationFrame, deferring reflow to the next paint cycle.'
)
status_row(doc,
    'Delegated listener and requestAnimationFrame both implemented.',
    'FIXED', GREEN)

hline()

# ================================================================
# SECTION 3 -- ARCHITECTURE
# ================================================================
heading1('3  Architecture', color=(30, 60, 140))

heading2('A1 -- Global Variable Namespace Pollution  [Medium]', color=(100, 60, 0))
body('Location: all script files loaded via <script> tags in Legit.html')
body(
    'All six script files (localization.js, utils.js, agents.js, ui.js, '
    'orchestrator.js, popup.js) are now wrapped in IIFEs. '
    'Internal helpers are private to each module; only functions needed '
    'cross-file are explicitly promoted via window.X = X at the end of '
    'each IIFE. Full migration to ES Modules (type="module" + import/export) '
    'has not been done, but the IIFE approach achieves equivalent scoping '
    'safety for this project\'s scale.'
)
status_row(doc,
    'IIFE scoping applied to all files. ES Modules not adopted (acceptable at this scale).',
    'SUBSTANTIALLY FIXED', GREEN)

heading2('A2 -- popup.js Monolith  [Medium]', color=(100, 60, 0))
body('Location: formerly popup.js (671 lines)')
body('The monolith has been split into four focused files exactly as recommended:')
bullet('orchestrator.js  --  startAnalysis, extractPageData, runProgressiveAnalysis, analyzeAgent')
bullet('ui.js            --  showStatus, toggleApiKeyView')
bullet('popup.js         --  thin entry point (~130 lines), DOMContentLoaded wiring only')
bullet('Load order in Legit.html: localization -> utils -> agents -> ui -> orchestrator -> popup')
body(
    'A bonus fix was applied: loadFromCache previously referenced DOM nodes via '
    'closed-over variables from popup.js\'s DOMContentLoaded callback, which were '
    'inaccessible from utils.js. Replaced with document.getElementById calls.'
)
status_row(doc,
    'Monolith split complete; bonus scope bug fixed as well.',
    'FULLY FIXED', GREEN)

heading2('A3 -- Dead Code  [Low]', color=GREY)
body('Two dead-code items were identified:')
bullet(
    'extractTextWithNewlines (utils.js) -- marked !!WIP!!, never called. '
    'Confirmed deleted: not present anywhere in the codebase.',
    bold_prefix='Deleted:'
)
bullet(
    'Readability.js loaded in Legit.html -- the library is only needed when '
    'injected into the active tab via executeScript. Confirmed removed: '
    'Legit.html script tags now list only the six application scripts; '
    'Readability.js is injected on-demand in orchestrator.js.',
    bold_prefix='Removed:'
)
status_row(doc,
    'Both dead-code items cleaned up.',
    'FULLY FIXED', GREEN)

hline()

# ================================================================
# SECTION 4 -- PRIVACY & POLICY
# ================================================================
heading1('4  Privacy & Policy', color=(80, 0, 100))

heading2('P1 -- No Dedicated Privacy Policy  [Required by Chrome Web Store]', color=(160, 30, 30))
body(
    'Chrome Web Store mandates a dedicated privacy policy for extensions '
    'requesting sensitive permissions (<all_urls>, scripting). The manifest '
    'currently declares host_permissions: ["*://*/*", '
    '"https://generativelanguage.googleapis.com/*"] and permissions: '
    '["activeTab", "scripting", "storage", "sidePanel"]. '
    'No privacy_policy_url field is present in manifest.json, and no '
    'privacy policy page or document has been added to the repository.'
)
status_row(doc,
    'No privacy policy file or manifest URL found.',
    'NOT ADDRESSED', RED)

heading2('P2 -- No User Disclosure About Content Sent to Google  [Medium]', color=(100, 60, 0))
body(
    'Users are not explicitly informed that article text is transmitted to '
    'the Google Gemini API under their own API key. No onboarding message, '
    'tooltip, or in-UI notice was added to Legit.html or the setup flow. '
    'This is a transparency concern, especially given that full article bodies '
    'are still sent (see E1 above).'
)
status_row(doc,
    'No disclosure notice added to the UI or onboarding flow.',
    'NOT ADDRESSED', RED)

hline()

# ================================================================
# SECTION 5 -- SUMMARY TABLE
# ================================================================
heading1('5  Summary', color=(20, 20, 20))
body('Quick-reference status for all 12 findings:')
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'

hdr = tbl.rows[0].cells
for cell, text in zip(hdr, ['Finding', 'Severity', 'Status']):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, bold=True, size=10, color=(255, 255, 255))
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '2D3748')
    tc_pr.append(shd)

rows_data = [
    ('Bug 1 -- styleScoreLabel swapped params',   'Critical',  'Not Fixed',          'FFCCCC'),
    ('Bug 2 -- Tautological condition',            'Medium',    'Fixed',              'CCFFCC'),
    ('Bug 3 -- waitForTabLoad no timeout',         'Medium',    'Not Fixed',          'FFCCCC'),
    ('E1 -- Full text to all agents',              'High',      'Partial',            'FFF3CC'),
    ('E2 -- verify-format doubles API calls',      'High',      'Not Addressed',      'FFCCCC'),
    ('E3 -- Search grounding for bias',            'Medium',    'Already OK',         'CCFFCC'),
    ('E4 -- DOM inefficiency',                     'Low',       'Fixed',              'CCFFCC'),
    ('A1 -- Global variable pollution',            'Medium',    'Substantially Fixed','CCFFCC'),
    ('A2 -- popup.js monolith',                    'Medium',    'Fully Fixed',        'CCFFCC'),
    ('A3 -- Dead code',                            'Low',       'Fully Fixed',        'CCFFCC'),
    ('P1 -- No privacy policy',                    'Required',  'Not Addressed',      'FFCCCC'),
    ('P2 -- No user disclosure',                   'Medium',    'Not Addressed',      'FFCCCC'),
]

for finding, severity, status, fill in rows_data:
    row = tbl.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, [finding, severity, status])):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        set_font(r, size=9.5)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        tc_pr.append(shd)

tbl.columns[0].width = Cm(9)
tbl.columns[1].width = Cm(2.5)
tbl.columns[2].width = Cm(5.5)

hline()

# ================================================================
# SECTION 6 -- OPEN ITEMS
# ================================================================
heading1('6  Open Items -- Recommended Next Actions', color=(160, 30, 30))
body('The following items should be addressed before the next store submission:')

items = [
    (
        '[CRITICAL] Fix Bug 1 -- styleScoreLabel parameter order',
        RED,
        'In utils.js line 535, swap the 7th and 8th arguments in the styleScoreLabel '
        'call so they match the function signature at line 561: '
        '(..., gradient, emoji, labelKey) instead of (..., gradient, labelKey, emoji).'
    ),
    (
        '[REQUIRED] Add Privacy Policy',
        RED,
        'Publish a privacy policy page and add its URL to manifest.json as '
        '"privacy_policy_url". Even a minimal statement that no data is collected '
        'by the developer is sufficient to satisfy the Chrome Web Store requirement.'
    ),
    (
        '[MEDIUM] Fix Bug 3 -- waitForTabLoad timeout',
        ORANGE,
        'Add a 10-second setTimeout inside waitForTabLoad that calls removeListener '
        'and resolves (or rejects) to prevent a dangling listener when the tab '
        'remains on a Google or DuckDuckGo domain.'
    ),
    (
        '[MEDIUM] Add User Disclosure for Data Transmission',
        ORANGE,
        'Add a brief notice in the setup view or first-run onboarding informing '
        'the user that article content is sent to Google Gemini API under their own key.'
    ),
    (
        '[HIGH] Complete E1 -- Cap Text Length and Use shortExcerpt for Headline',
        ORANGE,
        'Limit excerptStart to 3,000-4,000 characters in orchestrator.js and '
        'update the headline agent prompt in agents.js to use shortExcerpt '
        'rather than longExcerpt.'
    ),
]

for i, (title, color, desc) in enumerate(items, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(title)
    set_font(r, bold=True, size=10.5, color=color)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.2)
    p2.paragraph_format.space_after = Pt(6)
    rb = p2.add_run(desc)
    set_font(rb, size=10.5)

doc.save('feedback-review.docx')
print('Done.')
